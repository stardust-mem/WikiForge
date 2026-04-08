"""DOCX 文档提取 — python-docx"""

import io
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT


@dataclass
class DOCXExtractResult:
    text: str
    headings: list[dict]  # [{"level": 1, "text": "...", "char_offset": 0}]
    images: list[bytes]
    image_types: list[str]
    tables: list[str]  # markdown 格式的表格


def extract_docx(file_path: Path) -> DOCXExtractResult:
    """
    提取 DOCX 内容：段落文字 + Heading 层级 + 嵌入图片 + 表格

    输出：
    - text: 完整文本（保留 Heading 标记）
    - headings: Heading 位置列表（用于结构检测）
    - images: 图片 bytes 列表
    - tables: markdown 格式表格列表
    """
    doc = Document(str(file_path))

    text_parts = []
    headings = []
    images = []
    image_types = []
    tables = []
    char_offset = 0

    # 提取段落
    for para in doc.paragraphs:
        style_name = para.style.name if para.style else ""

        # 检测 Heading
        if style_name.startswith("Heading"):
            try:
                level = int(style_name.split()[-1])
            except (ValueError, IndexError):
                level = 1
            headings.append({
                "level": level,
                "text": para.text,
                "char_offset": char_offset,
            })
            prefix = "#" * level + " "
            line = prefix + para.text
        else:
            line = para.text

        text_parts.append(line)
        char_offset += len(line) + 1  # +1 for newline

        # 检测段落中的嵌入图片
        for run in para.runs:
            for drawing in run.element.findall(
                ".//{http://schemas.openxmlformats.org/drawingml/2006/main}blip"
            ):
                embed_attr = drawing.get(
                    "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
                )
                if embed_attr:
                    try:
                        rel = doc.part.rels[embed_attr]
                        img_bytes = rel.target_part.blob
                        if len(img_bytes) > 5000:
                            content_type = rel.target_part.content_type or "image/png"
                            images.append(img_bytes)
                            image_types.append(content_type)
                    except (KeyError, AttributeError):
                        continue

    # 提取表格为 markdown
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append("| " + " | ".join(cells) + " |")
        if rows:
            # 插入表头分隔行
            header = rows[0]
            sep = "| " + " | ".join(["---"] * len(table.rows[0].cells)) + " |"
            md_table = "\n".join([header, sep] + rows[1:])
            tables.append(md_table)

    text = "\n".join(text_parts)

    return DOCXExtractResult(
        text=text,
        headings=headings,
        images=images,
        image_types=image_types,
        tables=tables,
    )
